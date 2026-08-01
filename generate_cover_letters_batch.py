#!/usr/bin/env python3
"""Generate cover letters for jobs in batches (credit-efficient)."""

import sqlite3
import time
from datetime import datetime
from pathlib import Path
from config import Config
from utils import format_cv_label
from cover_letter_generator import CoverLetterGenerator
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def save_as_docx(text, filepath):
    """Save cover letter text as a formatted Word document."""
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if not line.strip():
            p.paragraph_format.space_after = Pt(0)

    doc.save(filepath)
    return filepath

def get_jobs_without_letters(limit=5):
    """Get jobs from DB that don't have cover letters yet, including best_cv."""
    try:
        conn = sqlite3.connect(Config.DATABASE_PATH)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id, job_title, company, description, link, best_cv
            FROM jobs
            WHERE id NOT IN (SELECT job_id FROM cover_letters_sent)
            ORDER BY relevance_score DESC, extracted_date DESC
            LIMIT ?
        """, (limit,))
        jobs = cursor.fetchall()
        conn.close()
        return jobs
    except Exception as e:
        print(f"[!] Error fetching jobs: {e}")
        return []

def mark_letter_generated(job_id):
    """Mark job as having a cover letter."""
    try:
        conn = sqlite3.connect(Config.DATABASE_PATH)
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO cover_letters_sent (job_id)
            VALUES (?)
            ON CONFLICT(job_id) DO NOTHING
        """, (job_id,))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        print(f"[!] Error marking letter as generated: {e}")
        return False

def send_telegram_status(generated_count):
    """Send status update to Telegram."""
    if not Config.TELEGRAM_BOT_TOKEN or not Config.TELEGRAM_CHAT_ID:
        return

    message = f"✉️ <b>Cover Letters Generated</b>\n"
    message += f"Generated <b>{generated_count}</b> new letter{'s' if generated_count != 1 else ''}\n"
    message += f"<i>{datetime.now().strftime('%Y-%m-%d %H:%M')}</i>"

    url = f"https://api.telegram.org/bot{Config.TELEGRAM_BOT_TOKEN}/sendMessage"
    payload = {'chat_id': Config.TELEGRAM_CHAT_ID, 'text': message, 'parse_mode': 'HTML'}

    try:
        requests.post(url, data=payload, timeout=10)
    except Exception as e:
        print(f"[!] Telegram notification failed: {e}")

def main():
    print("=" * 70)
    print("BATCH COVER LETTER GENERATOR (Credit-Efficient)")
    print("=" * 70)

    # Get jobs needing letters (limit to 3 per run, optimized for credits)
    jobs = get_jobs_without_letters(limit=3)

    if not jobs:
        print("[*] No jobs need cover letters right now.")
        return 0

    print(f"[*] Found {len(jobs)} job(s) needing cover letters")
    print()

    generator = CoverLetterGenerator()
    generated_count = 0

    for job_id, title, company, description, link, best_cv in jobs:
        cv_label = format_cv_label(best_cv)
        print(f"[*] Generating letter for: {title} @ {company}  [CV: {cv_label or 'auto'}]")

        try:
            # Generate letter using the pre-scored best CV
            letter = generator.generate_cover_letter(title, company, description,
                                                      cv_name_hint=best_cv)

            if letter:
                # Save to file
                safe_title = title.replace('/', '-').replace('\\', '-')[:50]
                docx_dir = Path(Config.OUTPUT_DIR) / "docx cover letters"
                docx_dir.mkdir(exist_ok=True)
                filename = str(docx_dir / f"cover_letter_{job_id}_{safe_title}.docx")
                save_as_docx(letter, filename)

                # Mark as generated
                mark_letter_generated(job_id)

                print(f"    [+] Letter saved to {filename}")
                generated_count += 1

                # Small delay between API calls to avoid rate limiting
                print(f"    [*] Waiting 2 seconds before next letter...")
                time.sleep(2)
            else:
                print(f"    [!] Failed to generate letter")

        except Exception as e:
            print(f"    [!] Error: {e}")
            # Still wait to avoid hammering API
            time.sleep(2)

    print()
    print("=" * 70)
    print(f"[+] Generated {generated_count} new cover letter(s)")
    print("=" * 70)

    # Notify via Telegram if any were generated
    if generated_count > 0:
        send_telegram_status(generated_count)

    return generated_count

if __name__ == '__main__':
    main()
