"""
Interactive path: search, then draft cover letters for the top few matches.

This is not the scheduled pipeline and it does not deliver anything. The
daily run is job_search_smart.py to search and store, then telegram_sender.py
to compose and send the digest. Use this script by hand when you want DOCX
cover letters written for the best current matches.
"""

import sys
import json
from pathlib import Path
from job_search_smart import SmartJobSearcher
from cover_letter_generator import CoverLetterGenerator
from database import Database
from utils import setup_logging
from config import Config
from docx import Document
from docx.shared import Pt, Inches


def save_as_docx(text, filepath):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if not line.strip():
            p.paragraph_format.space_after = Pt(0)
    doc.save(filepath)
    return filepath

logger = setup_logging('run')

def print_header(text):
    """Print a formatted header."""
    print("\n" + "="*70)
    print(f"  {text}")
    print("="*70)

def run_job_search():
    """Search for remote jobs and rank by relevance."""
    print_header("SEARCHING FOR REMOTE JOBS")

    searcher = SmartJobSearcher()
    jobs = searcher.search_all_sources()

    if not jobs:
        print("[-] No jobs found!")
        return []

    print(f"\n[+] Found {len(jobs)} relevant remote jobs:\n")
    for i, job in enumerate(jobs[:10], 1):
        print(f"{i}. {job['title']} @ {job['company']}")
        print(f"   Relevance: {job.get('relevance_score', 0)}% | {job['source']}")
        print(f"   Link: {job.get('link', 'N/A')[:60]}...")
        print()

    return jobs

def generate_cover_letters(jobs, num_to_generate=3):
    """Generate cover letters for top matching jobs."""
    print_header(f"GENERATING COVER LETTERS FOR TOP {num_to_generate} JOBS")

    generator = CoverLetterGenerator()
    db = Database()
    cover_letters = []

    for i, job in enumerate(jobs[:num_to_generate], 1):
        print(f"\n[*] Generating cover letter {i}/{num_to_generate}...")
        print(f"    {job['title']} @ {job['company']}")

        letter = generator.generate_cover_letter(
            job.get('title'),
            job.get('company'),
            job.get('description'),
            cv_name_hint=job.get('best_cv')
        )

        if letter:
            formatted = generator.format_cover_letter(
                job.get('title'),
                job.get('company'),
                letter
            )

            cover_letters.append({
                'job_id': job.get('id'),
                'job_title': job.get('title'),
                'company': job.get('company'),
                'cover_letter': formatted
            })

            # Save to database
            db.add_cover_letter(
                job_id=job.get('id'),
                job_title=job.get('title'),
                company=job.get('company'),
                selected_cv='Auto-Selected',
                generated_letter=formatted
            )

            print(f"    [+] Cover letter generated")

    return cover_letters

def save_cover_letters(cover_letters):
    """Save cover letters to output directory."""
    output_dir = Config.OUTPUT_DIR
    output_dir.mkdir(exist_ok=True)

    saved_files = []

    for cl in cover_letters:
        docx_dir = output_dir / "docx cover letters"
        docx_dir.mkdir(exist_ok=True)
        filename = f"Cover_Letter_{cl['company'].replace(' ', '_')}_{cl['job_title'].replace(' ', '_')}.docx"
        filepath = docx_dir / filename

        try:
            save_as_docx(cl['cover_letter'], filepath)
            saved_files.append(filepath)
            print(f"[+] Saved: {filename}")
        except Exception as e:
            logger.error(f"Error saving cover letter: {e}")

    return saved_files

def main():
    """Run the complete job search and cover letter system."""
    print("\n" + "="*70)
    print("  JOB SEARCH + COVER LETTER GENERATION SYSTEM")
    print("="*70)

    try:
        # Step 1: Search for jobs
        jobs = run_job_search()
        if not jobs:
            print("\n[-] No jobs to process. Exiting.")
            return 1

        # Step 2: Generate cover letters
        cover_letters = generate_cover_letters(jobs, num_to_generate=min(3, len(jobs)))

        if not cover_letters:
            print("\n[-] No cover letters generated. Exiting.")
            return 1

        # Step 3: Save cover letters
        print_header("SAVING COVER LETTERS")
        saved = save_cover_letters(cover_letters)
        print(f"\n[+] Saved {len(saved)} cover letters to {Config.OUTPUT_DIR}")

        # Summary
        print_header("SUMMARY")
        print(f"[+] Jobs found: {len(jobs)}")
        print(f"[+] Cover letters generated: {len(cover_letters)}")
        print(f"[+] Cover letters saved: {len(saved)}")
        print(f"\n[+] Cover letters are in: {Config.OUTPUT_DIR}")
        print("\nNext steps:")
        print("1. Review the cover letters")
        print("2. Customize if needed")
        print("3. Send to companies with your CV")

        logger.info(f"Job search completed: {len(jobs)} jobs, {len(cover_letters)} cover letters")
        return 0

    except Exception as e:
        logger.error(f"Error in main execution: {e}")
        print(f"\n[-] Error: {e}")
        return 1

if __name__ == '__main__':
    sys.exit(main())
